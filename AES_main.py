from docx import Document
import  AES_128_MaHoa as MaHoaTool
import  AES_128_GiaiMa as GiaiMaTool
from tkinter import filedialog

def open_file_dialog():
    file_path = filedialog.askopenfilename()
    print("Đường dẫn file đã chọn:", file_path)
    return  file_path

def Read_File_Word(file_path):
    try:
        doc = Document(file_path)
        content = []

        for paragraph in doc.paragraphs:
            content.append(paragraph.text)

        return ''.join(content)
    except FileNotFoundError:
        print(f"File not found: {file_path}")
    except Exception as e:
        print(f"Error: {e}")

def Chose_File_Word():
    file_path = open_file_dialog()
    text = Read_File_Word(file_path)

    return text

def Save_To_Word_File( content, file_path):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)

def MaHoa( fp_ban_ro, fb_key, fb_ban_mat ):

    ban_ro = Read_File_Word(fp_ban_ro)
    key = Read_File_Word(fb_key)

    ban_mat = MaHoaTool.MaHoa_Text(ban_ro,key)
    ban_mat = "" + ban_mat
    Save_To_Word_File( ban_mat , fb_ban_mat)

def MaHoa_02( ban_ro, key ):
    ban_mat = MaHoaTool.MaHoa_Text(ban_ro, key)
    ban_mat = "" + ban_mat

    Save_To_Word_File(ban_mat, 'BanMat.docx')

    return  ban_mat

def GiaiMa( fp_ban_mat, fp_key ):

    ban_mat = Read_File_Word(fp_ban_mat)
    key = Read_File_Word(fp_key)

    ban_ro = GiaiMaTool.GiaiMa_Text(ban_mat, key)

    return  ban_ro

def GiaiMa_02( ban_mat, key ):

    ban_ro = GiaiMaTool.GiaiMa_Text(ban_mat, key)

    return  ban_ro

# fp_banRo = 'BanRo.docx'
# fp_banRo = open_file_dialog()
# fp_banMat = 'BanMat.docx'
# fp_key = 'Key.docx'
#
# MaHoa( fp_banRo, fp_key, fp_banMat)
# print("Ban ro: " , GiaiMa(fp_banMat, fp_key))


